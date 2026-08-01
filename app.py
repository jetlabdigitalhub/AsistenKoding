import os
from flask import Flask, request, jsonify, send_file, send_from_directory, render_template_string, g
from workspace.highlight_engine import HighlightEngine
from workspace.memo_engine import MemoEngine
from workspace.export_engine import ExportEngine
from workspace.semantic_engine import SemanticEngine
from workspace.semantic_annotation import detect_indicators, suggest_codes, generate_semantic_relationship
from coding_engine.pipeline import process_transcript as pipeline_process
from backend.codeweaving import generate as generate_codeweaving
from backend.evidence_matcher import find_supporting_chunks, find_disconfirming_chunks
from backend.session import get_session_id
import io
from docx import Document as DocxDocument
import json
from pathlib import Path
from datetime import datetime
from sumy.parsers.plaintext import PlaintextParser
from sumy.nlp.tokenizers import Tokenizer
from sumy.summarizers.lsa import LsaSummarizer

BASE_DIR = os.path.dirname(__file__)

app = Flask(__name__, static_folder="frontend", static_url_path="", template_folder="frontend")

highlight = HighlightEngine(storage_dir=os.path.join(BASE_DIR, "data"))
memo = MemoEngine(storage_dir=os.path.join(BASE_DIR, "data"))
exporter = ExportEngine()
semantic = SemanticEngine(modules_dir=os.path.join(BASE_DIR, "modules"))


@app.before_request
def require_session_for_api():
    if request.path.startswith('/api/'):
        session_id = get_session_id()
        if not session_id:
            print(f"[Session] Missing X-Session-ID from {request.remote_addr} for {request.path}")
            return jsonify({'error': 'Missing or invalid X-Session-ID header'}), 400
        g.session_id = session_id
        try:
            print(f"[Session] {request.remote_addr} -> {g.session_id} {request.path}")
        except Exception:
            pass


def _data_file(kind, doc_id='default'):
    session_id = getattr(g, 'session_id', None)
    if not session_id:
        raise RuntimeError('session_id not available for data persistence')
    session_dir = Path(os.path.join(BASE_DIR, 'data', 'sessions', session_id))
    session_dir.mkdir(parents=True, exist_ok=True)
    return session_dir / f"{kind}_{doc_id}.json"


def _read_json(kind, doc_id='default'):
    p = _data_file(kind, doc_id)
    if not p.exists():
        return []
    try:
        return json.loads(p.read_text(encoding='utf-8'))
    except Exception:
        return []


def _write_json(kind, doc_id, data):
    p = _data_file(kind, doc_id)
    p.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding='utf-8')
    return True


@app.route('/')
def index():
    return send_from_directory('frontend', 'index.html')


@app.route('/api/modules')
def list_modules():
    semantic.load_all()
    return jsonify({"modules": semantic.list_modules()})


@app.route('/api/analyze', methods=['POST'])
def analyze():
    data = request.json or {}
    mod = data.get('module')
    text = data.get('text', '')
    if not mod:
        return jsonify({"error": "module required"}), 400
    return jsonify(semantic.analyze_with(mod, text))


@app.route('/api/pipeline', methods=['POST'])
def pipeline_api():
    data = request.json or {}
    text = data.get('text', '')
    # run local explainable pipeline
    try:
        res = pipeline_process(text)
        return jsonify(res)
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/highlight', methods=['POST'])
def add_highlight():
    data = request.json or {}
    doc_id = data.get('doc_id', 'default')
    session_id = g.session_id
    # create highlight first
    h = highlight.save_highlight(
        session_id,
        doc_id,
        data.get('start'),
        data.get('end'),
        data.get('text'),
        data.get('code') or data.get('label'),
        data.get('module'),
        category=data.get('category')
    )
    # run linguistic detection automatically on created highlight text only if user didn't provide codes/semantic
    try:
        if data.get('codes'):
            # user provided selected codes; respect them and optionally semantic_relationship
            allowed = {}
            if data.get('codes'):
                allowed['codes'] = data.get('codes')
            if data.get('semantic_relationship'):
                allowed['semantic_relationship'] = data.get('semantic_relationship')
            if data.get('color'):
                allowed['color'] = data.get('color')
            if allowed:
                h = highlight.update_highlight(session_id, doc_id, h['id'], **allowed)
        else:
            indicators = detect_indicators(h.get('text',''))
            suggestions = suggest_codes(indicators)
            semantic_rel = generate_semantic_relationship(indicators, h.get('text',''))
            # extract ranked codes (strings)
            ranked = [s['code'] for s in suggestions]
            # update highlight with codes and semantic relationship
            h = highlight.update_highlight(session_id, doc_id, h['id'], codes=ranked, semantic_relationship=semantic_rel)
    except Exception:
        pass
    return jsonify(h)


@app.route('/api/highlight/<hid>', methods=['PUT'])
def edit_highlight(hid):
    data = request.json or {}
    doc_id = data.get('doc_id', 'default')
    session_id = g.session_id
    allowed = {k: v for k, v in data.items() if k in ['text','start','end','code','cycle','memo','module','category','semantic','color','codes','semantic_relationship']}
    h = highlight.update_highlight(session_id, doc_id, hid, **allowed)
    if not h:
        return jsonify({"error":"not found"}), 404
    return jsonify(h)


@app.route('/api/highlight/<hid>', methods=['DELETE'])
def remove_highlight(hid):
    doc_id = request.args.get('doc_id', 'default')
    session_id = g.session_id
    ok = highlight.delete_highlight(session_id, doc_id, hid)
    return jsonify({"deleted": ok})


@app.route('/api/highlights')
def get_highlights():
    doc_id = request.args.get('doc_id', 'default')
    session_id = g.session_id
    return jsonify(highlight.list_highlights(session_id, doc_id))


@app.route('/api/clear_highlights', methods=['POST'])
def clear_highlights():
    data = request.json or {}
    doc_id = data.get('doc_id', 'default')
    session_id = g.session_id
    highlight.clear_highlights(session_id, doc_id)
    return jsonify({"cleared": True})


@app.route('/api/memo', methods=['POST'])
def add_memo():
    data = request.json or {}
    doc_id = data.get('doc_id', 'default')
    session_id = g.session_id
    entry = memo.add_memo(session_id, doc_id, data.get('author', 'anon'), data.get('text', ''))
    return jsonify(entry)


@app.route('/api/memos')
def list_memos():
    doc_id = request.args.get('doc_id', 'default')
    session_id = g.session_id
    return jsonify(memo.list_memos(session_id, doc_id))


@app.route('/api/upload_docx', methods=['POST'])
def upload_docx():
    if 'file' not in request.files:
        return jsonify({"error": "no file"}), 400
    f = request.files['file']
    try:
        fc = _read_json('first_cycle', 'default')
        print(f"[Upload] Incoming upload — current first_cycle count for project default: {len(fc)}")
    except Exception:
        print('[Upload] Could not read first_cycle before upload')
    try:
        data = f.read()
        bio = io.BytesIO(data)
        doc = DocxDocument(bio)
        # collect paragraph text
        paragraphs = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
        # also collect text from tables (common in exported docs)
        try:
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell_text = (cell.text or '').strip()
                        if cell_text:
                            paragraphs.append(cell_text)
        except Exception:
            # non-fatal if tables can't be iterated
            pass
        text = '\n\n'.join(paragraphs).strip()
        # fallback: if no paragraphs found, try to decode raw bytes as text
        if not text:
            try:
                decoded = data.decode('utf-8')
            except Exception:
                try:
                    decoded = data.decode('latin-1')
                except Exception:
                    decoded = ''
            text = decoded.strip()
        print(f"[Upload] Parsed text length: {len(text)}")
        # store upload metadata so exports can reference filename and timestamp
        try:
            _write_json('upload_info', 'default', {'filename': getattr(f, 'filename', None) or 'N/A', 'uploaded_at': datetime.utcnow().isoformat()})
        except Exception as e:
            print('[Upload] Failed to write upload_info', e)
        return jsonify({"text": text, 'filename': getattr(f, 'filename', None)})
    except Exception as e:
        print(f"[Upload] Error parsing .docx: {e}")
        return jsonify({"error": str(e)}), 500


@app.route('/api/detect', methods=['POST'])
def detect():
    data = request.json or {}
    text = data.get('text', '')
    try:
        indicators = detect_indicators(text)
        suggestions = suggest_codes(indicators)
        semantic_rel = generate_semantic_relationship(indicators, text)
        return jsonify({"indicators": indicators, "suggestions": suggestions, "semantic_relationship": semantic_rel})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route('/api/summarize', methods=['POST'])
def summarize_api():
    data = request.json or {}
    text = data.get('text', '')
    sentences = int(data.get('sentences', 1))
    language = data.get('language', 'english')
    if not text:
        return jsonify({'summary': []})
    try:
        try:
            parser = PlaintextParser.from_string(text, Tokenizer(language))
            summarizer = LsaSummarizer()
            summary = summarizer(parser.document, sentences)
            sents = [str(s).strip() for s in summary]
            return jsonify({'summary': sents})
        except Exception:
            # Fallback: lightweight frequency-based summarizer (no external resources)
            import re
            sents = re.split(r'(?<=[.!?])\s+', text.strip())
            if not sents:
                return jsonify({'summary': []})
            # simple tokenization and frequency
            words = re.findall(r"\w+", text.lower())
            stopwords = set(['the','and','is','in','to','of','a','for','that','it','on','as','are','with','this','by','an','be','or','from','at','was','which','but','have','has'])
            freq = {}
            for w in words:
                if w in stopwords or len(w) <= 2:
                    continue
                freq[w] = freq.get(w, 0) + 1
            # score sentences
            scored = []
            for i, s in enumerate(sents):
                ws = re.findall(r"\w+", s.lower())
                score = sum(freq.get(w,0) for w in ws)
                scored.append((i, score, s))
            scored.sort(key=lambda x: x[1], reverse=True)
            top = sorted(scored[:max(1, min(len(scored), sentences))], key=lambda x: x[0])
            out = [t[2].strip() for t in top]
            return jsonify({'summary': out})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/export', methods=['POST'])
def do_export():
    data = request.json or {}
    doc_id = data.get('doc_id', 'default')
    session_id = g.session_id
    text = data.get('text', '')
    highlights = highlight.list_highlights(session_id, doc_id)
    memos = memo.list_memos(session_id, doc_id)
    # load first_cycle, second_cycle, rich memos
    first_cycle = _read_json('first_cycle', doc_id)
    second_cycle = _read_json('second_cycle', doc_id)
    memos_rich = _read_json('memos_rich', doc_id)
    # collect analyses per module present in highlights
    modules_in_highlights = sorted(list({h.get('module') for h in highlights if h.get('module')}))
    module_analyses = []
    for mname in modules_in_highlights:
        try:
            module_analyses.append(semantic.analyze_with(mname, text))
        except Exception:
            module_analyses.append({"module": {"name": mname}, "suggestions": [], "analysis": {}})
    # include first/second cycle summaries
    module_analyses.append({'module': {'name': 'first_cycle'}, 'suggestions': [], 'analysis': {'first_cycle': first_cycle}})
    module_analyses.append({'module': {'name': 'second_cycle'}, 'suggestions': [], 'analysis': {'second_cycle': second_cycle}})
    module_analyses.append({'module': {'name': 'memos_rich'}, 'suggestions': [], 'analysis': {'memos': memos_rich}})
    # run pipeline to obtain chunk segmentation for export
    try:
        pipeline_result = pipeline_process(text)
        chunks_for_export = pipeline_result.get('chunks', []) if isinstance(pipeline_result, dict) else []
    except Exception:
        pipeline_result = None
        chunks_for_export = []
    fmt = data.get('format', 'docx')
    codeweaving_items = _read_json('codeweaving_items', doc_id)
    uploaded_filename = data.get('uploaded_filename')
    uploaded_at = data.get('uploaded_at')
    # merge memos_rich into simple memos list for exporter compatibility
    merged_memos = []
    try:
        merged_memos = list(memos or [])
        for mr in (memos_rich or []):
            # normalize rich memo to a simple memo dict with text and chunk_id/created/author
            merged_memos.append({
                'text': mr.get('memo_text') or mr.get('text') or '',
                'chunk_id': mr.get('chunk_id') or mr.get('chunk') or None,
                'author': mr.get('author'),
                'created': mr.get('created') or mr.get('created_at')
            })
    except Exception:
        merged_memos = memos or []

    session_id = g.session_id
    if fmt == 'pdf':
        # if upload metadata not provided in request, try stored upload_info
        if not uploaded_filename or not uploaded_at:
            try:
                info = _read_json('upload_info', doc_id)
                if isinstance(info, dict):
                    if not uploaded_filename:
                        uploaded_filename = info.get('filename')
                    if not uploaded_at:
                        uploaded_at = info.get('uploaded_at')
            except Exception:
                pass
        try:
            path = exporter.export_pdf(doc_id, text, highlights=highlights, memos=merged_memos, module_analyses=module_analyses, first_cycle=first_cycle, second_cycle=second_cycle, codeweaving_items=codeweaving_items, chunks=chunks_for_export, uploaded_filename=uploaded_filename, uploaded_at=uploaded_at, session_id=session_id)
        except Exception as e:
            print('[Export] PDF export error:', e)
            return jsonify({'error': str(e)}), 500
        # return the generated PDF file directly to the browser for download
        filename = os.path.basename(path)
        return send_file(path, mimetype='application/pdf', as_attachment=True, download_name=filename)
    else:
        if not uploaded_filename or not uploaded_at:
            try:
                info = _read_json('upload_info', doc_id)
                if isinstance(info, dict):
                    if not uploaded_filename:
                        uploaded_filename = info.get('filename')
                    if not uploaded_at:
                        uploaded_at = info.get('uploaded_at')
            except Exception:
                pass
        try:
            path = exporter.export_docx(doc_id, text, highlights=highlights, memos=merged_memos, module_analyses=module_analyses, first_cycle=first_cycle, second_cycle=second_cycle, codeweaving_items=codeweaving_items, chunks=chunks_for_export, uploaded_filename=uploaded_filename, uploaded_at=uploaded_at, session_id=session_id)
        except Exception as e:
            print('[Export] DOCX export error:', e)
            return jsonify({'error': str(e)}), 500
    return jsonify({"path": path})


@app.route('/api/first_cycle', methods=['GET'])
def list_first_cycle():
    doc_id = request.args.get('doc_id', 'default')
    return jsonify(_read_json('first_cycle', doc_id))


@app.route('/api/first_cycle', methods=['POST'])
def add_first_cycle():
    data = request.json or {}
    doc_id = data.get('doc_id', 'default')
    entries = _read_json('first_cycle', doc_id)
    chunk_id = data.get('chunk_id')
    code = data.get('code') or data.get('label')
    now = datetime.utcnow().isoformat()
    source = data.get('source_type') or ('manual' if data.get('selected_by_user') else 'ai')
    color = data.get('color') or data.get('code_color') or None
    # deduplicate by chunk+code_name
    existing = None
    for e in entries:
        if e.get('chunk_id') == chunk_id and e.get('code_name') == code:
            existing = e
            break
    if existing:
        existing['count'] = existing.get('count', 1) + 1
        existing['updated_at'] = now
        print(f"[FirstCycle] Incremented existing code {code} for chunk {chunk_id}")
        saved = existing
    else:
        new = {
            'id': str(len(entries)+1),
            'project_id': doc_id,
            'chunk_id': chunk_id,
            'code_name': code,
            'code_color': color,
            'source_type': source,
            'indicator': data.get('indicator'),
            'speaker': data.get('speaker'),
            'semantic': data.get('semantic'),
            'created_at': now,
            'updated_at': now,
            'count': 1
        }
        entries.append(new)
        saved = new
        print(f"[FirstCycle] Inserted new code {code} for chunk {chunk_id}")
    try:
        _write_json('first_cycle', doc_id, entries)
        print(f"[FirstCycle] Saved first_cycle to disk for project {doc_id}")
    except Exception as e:
        print(f"[FirstCycle] Failed to write first_cycle: {e}")
        return jsonify({'error': str(e)}), 500
    # return the created/updated entry and full list for compatibility
    return jsonify({'entry': saved, 'all': entries})


@app.route('/api/first_cycle/<fid>', methods=['PUT'])
def update_first_cycle(fid):
    data = request.json or {}
    doc_id = data.get('doc_id', 'default')
    entries = _read_json('first_cycle', doc_id)
    found = None
    for e in entries:
        if str(e.get('id')) == str(fid):
            found = e
            break
    if not found:
        return jsonify({'error': 'not found'}), 404
    # allow renaming code and updating metadata (use canonical field names)
    now = datetime.utcnow().isoformat()
    if 'code_name' in data:
        found['code_name'] = data['code_name']
    if 'code_color' in data:
        found['code_color'] = data['code_color']
    for k in ['indicator', 'speaker', 'semantic', 'count', 'source_type']:
        if k in data:
            found[k] = data[k]
    found['updated_at'] = now
    _write_json('first_cycle', doc_id, entries)
    print(f"[FirstCycle] Updated entry {fid}")
    return jsonify(found)


@app.route('/api/first_cycle/<fid>', methods=['DELETE'])
def delete_first_cycle(fid):
    doc_id = request.args.get('doc_id', 'default')
    entries = _read_json('first_cycle', doc_id)
    new = [e for e in entries if str(e.get('id')) != str(fid)]
    _write_json('first_cycle', doc_id, new)
    return jsonify({'deleted': True})


@app.route('/api/first_cycle/merge', methods=['POST'])
def merge_first_cycle():
    data = request.json or {}
    doc_id = data.get('doc_id', 'default')
    target = data.get('target')
    sources = data.get('sources') or []
    if not target or not sources:
        return jsonify({'error': 'target and sources required'}), 400
    entries = _read_json('first_cycle', doc_id)
    # replace source codes with target (use code_name)
    for e in entries:
        if e.get('code_name') in sources:
            e['code_name'] = target
    # optionally dedupe: combine counts for same chunk+code
    combined = []
    keymap = {}
    for e in entries:
        key = (e.get('chunk_id'), e.get('code_name'))
        if key in keymap:
            keymap[key]['count'] = keymap[key].get('count',1) + e.get('count',1)
        else:
            keymap[key] = dict(e)
    combined = list(keymap.values())
    _write_json('first_cycle', doc_id, combined)
    return jsonify(combined)


@app.route('/api/codeweaving/generate', methods=['POST'])
def api_generate_codeweaving():
    data = request.json or {}
    session_id = g.session_id
    doc_id = data.get('doc_id', 'default')
    selected = data.get('selected_codes') or []
    mode = data.get('mode', 'auto')
    # run deterministic codeweaving
    try:
        assertion = generate_codeweaving(selected, mode=mode)
        highlights = highlight.list_highlights(session_id, doc_id)
        supporting = find_supporting_chunks(assertion.get('codes_used', []), highlights)
        disconfirm = find_disconfirming_chunks(assertion, highlights)
        return jsonify({'assertion': assertion, 'supporting': supporting, 'disconfirming': disconfirm})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/codeweaving/save', methods=['POST'])
def api_save_assertion():
    data = request.json or {}
    doc_id = data.get('doc_id', 'default')
    assertion = data.get('assertion')
    if not assertion:
        return jsonify({'error': 'assertion required'}), 400
    entries = _read_json('second_cycle', doc_id)
    # create a second-cycle entry storing assertion and trace
    new = {
        'id': str(len(entries)+1),
        'theme': data.get('theme') or f"theme_{len(entries)+1}",
        'assertion': assertion,
        'codes': assertion.get('codes_used', []),
        'chunks': assertion.get('source_chunks', []),
        'created': data.get('created')
    }
    entries.append(new)
    _write_json('second_cycle', doc_id, entries)
    return jsonify(new)


@app.route('/api/second_cycle', methods=['GET'])
def list_second_cycle():
    doc_id = request.args.get('doc_id', 'default')
    return jsonify(_read_json('second_cycle', doc_id))


@app.route('/api/second_cycle', methods=['POST'])
def create_second_cycle():
    data = request.json or {}
    doc_id = data.get('doc_id', 'default')
    entries = _read_json('second_cycle', doc_id)
    new = {
        'id': str(len(entries)+1),
        'theme': data.get('theme') or f"theme_{len(entries)+1}",
        'codes': data.get('codes') or [],
        'chunks': data.get('chunks') or [],
        'created': data.get('created')
    }
    entries.append(new)
    _write_json('second_cycle', doc_id, entries)
    return jsonify(new)


@app.route('/api/second_cycle/<tid>', methods=['PUT'])
def update_second_cycle(tid):
    data = request.json or {}
    doc_id = data.get('doc_id', 'default')
    entries = _read_json('second_cycle', doc_id)
    found = None
    for e in entries:
        if str(e.get('id')) == str(tid):
            found = e
            break
    if not found:
        return jsonify({'error':'not found'}), 404
    for k in ['theme','codes','chunks']:
        if k in data:
            found[k] = data[k]
    _write_json('second_cycle', doc_id, entries)
    return jsonify(found)


@app.route('/api/second_cycle/<tid>', methods=['DELETE'])
def delete_second_cycle(tid):
    doc_id = request.args.get('doc_id', 'default')
    entries = _read_json('second_cycle', doc_id)
    new = [e for e in entries if str(e.get('id')) != str(tid)]
    _write_json('second_cycle', doc_id, new)
    return jsonify({'deleted': True})


@app.route('/api/second_cycle/promote', methods=['POST'])
def promote_second_cycle():
    data = request.json or {}
    doc_id = data.get('doc_id', 'default')
    entries = _read_json('second_cycle', doc_id)
    # create or update theme by name
    theme = data.get('theme') or f"theme_{len(entries)+1}"
    existing = None
    for e in entries:
        if e.get('theme') == theme:
            existing = e
            break
    if existing:
        # append codes/chunks
        existing['codes'] = list(dict.fromkeys(existing.get('codes', []) + (data.get('codes') or [])))
        existing['chunks'] = list(dict.fromkeys(existing.get('chunks', []) + (data.get('chunks') or [])))
    else:
        new = {
            'id': str(len(entries)+1),
            'theme': theme,
            'codes': data.get('codes') or [],
            'chunks': data.get('chunks') or [],
            'created': data.get('created')
        }
        entries.append(new)
    _write_json('second_cycle', doc_id, entries)
    return jsonify(entries)


@app.route('/api/codeweaving_items', methods=['GET'])
def list_codeweaving_items():
    doc_id = request.args.get('doc_id', 'default')
    return jsonify(_read_json('codeweaving_items', doc_id))


@app.route('/api/codeweaving_items', methods=['POST'])
def create_codeweaving_item():
    data = request.json or {}
    doc_id = data.get('doc_id', 'default')
    entries = _read_json('codeweaving_items', doc_id)
    new = {
        'id': str(len(entries)+1),
        'narrative': data.get('narrative') or data.get('assertion_text') or '',
        'codes': data.get('codes') or data.get('codes_used') or [],
        'status': data.get('status') or 'suggested',
        'created': data.get('created'),
        'updated': data.get('created')
    }
    entries.append(new)
    _write_json('codeweaving_items', doc_id, entries)
    return jsonify(new)


@app.route('/api/codeweaving_items/<iid>', methods=['PUT'])
def update_codeweaving_item(iid):
    data = request.json or {}
    doc_id = data.get('doc_id', 'default')
    entries = _read_json('codeweaving_items', doc_id)
    found = None
    for e in entries:
        if str(e.get('id')) == str(iid):
            found = e
            break
    if not found:
        return jsonify({'error': 'not found'}), 404
    for k in ['narrative', 'codes', 'status']:
        if k in data:
            found[k] = data[k]
    from datetime import datetime
    found['updated'] = datetime.utcnow().isoformat()
    _write_json('codeweaving_items', doc_id, entries)
    return jsonify(found)


@app.route('/api/codeweaving_items/<iid>', methods=['DELETE'])
def delete_codeweaving_item(iid):
    doc_id = request.args.get('doc_id', 'default')
    entries = _read_json('codeweaving_items', doc_id)
    new = [e for e in entries if str(e.get('id')) != str(iid)]
    _write_json('codeweaving_items', doc_id, new)
    return jsonify({'deleted': True})


@app.route('/api/memos_rich', methods=['GET'])
def list_memos_rich():
    doc_id = request.args.get('doc_id', 'default')
    return jsonify(_read_json('memos_rich', doc_id))


@app.route('/api/memos_rich', methods=['POST'])
def add_memo_rich():
    data = request.json or {}
    doc_id = data.get('doc_id', 'default')
    entries = _read_json('memos_rich', doc_id)
    new = {
        'id': str(len(entries)+1),
        'author': data.get('author', 'anon'),
        'memo_text': data.get('memo_text') or data.get('text') or '',
        'chunk_id': data.get('chunk_id'),
        'code': data.get('code'),
        'cluster_id': data.get('cluster_id'),
        'created': data.get('created')
    }
    entries.append(new)
    _write_json('memos_rich', doc_id, entries)
    return jsonify(new)


@app.route('/api/memos_rich/<mid>', methods=['PUT'])
def update_memo_rich(mid):
    data = request.json or {}
    doc_id = data.get('doc_id', 'default')
    entries = _read_json('memos_rich', doc_id)
    found = None
    for e in entries:
        if str(e.get('id')) == str(mid):
            found = e
            break
    if not found:
        return jsonify({'error':'not found'}), 404
    for k in ['memo_text','chunk_id','code','cluster_id','author']:
        if k in data:
            found[k] = data[k]
    _write_json('memos_rich', doc_id, entries)
    return jsonify(found)


@app.route('/api/memos_rich/<mid>', methods=['DELETE'])
def delete_memo_rich(mid):
    doc_id = request.args.get('doc_id', 'default')
    entries = _read_json('memos_rich', doc_id)
    new = [e for e in entries if str(e.get('id')) != str(mid)]
    _write_json('memos_rich', doc_id, new)
    return jsonify({'deleted': True})


@app.route('/favicon.ico')
def favicon():
    # Serve favicon if present in frontend folder, otherwise return empty 204 to avoid 404 noise
    fav_path = os.path.join(BASE_DIR, 'frontend', 'favicon.ico')
    if os.path.exists(fav_path):
        return send_from_directory(os.path.join(BASE_DIR, 'frontend'), 'favicon.ico')
    return ('', 204)


if __name__ == '__main__':
    app.run(debug=True, port=5000)
