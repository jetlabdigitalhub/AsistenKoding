from pathlib import Path
from docx import Document, styles
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
import html
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
from reportlab.lib.units import inch
import datetime


class ExportEngine:
    def __init__(self, storage_dir=None):
        # default to user's Downloads folder when not provided
        if storage_dir:
            self.storage_dir = Path(storage_dir)
        else:
            self.storage_dir = Path.home() / 'Downloads'
        self.storage_dir.mkdir(parents=True, exist_ok=True)

    def export_docx(self, doc_id, text, highlights=None, memos=None, module_analyses=None, first_cycle=None, second_cycle=None, codeweaving_items=None, chunks=None, uploaded_filename=None, uploaded_at=None, out_name=None):
        # determine output filename: prefer explicit out_name, then uploaded_filename, then doc_id
        if out_name:
            out_name_final = out_name
        else:
            # choose extension
            ext = 'docx'
            name_base = None
            if uploaded_filename:
                # strip extension if present
                name_base = Path(uploaded_filename).stem
            elif doc_id and doc_id != 'default':
                name_base = str(doc_id)
            else:
                name_base = 'export'
            out_name_final = f"Coded - {name_base}.{ext}"
        path = self.storage_dir / out_name_final
        doc = Document()

        

        # Build chunk-level table and maps (collect transcripts, codes, second-cycle narrations, and memos)
        chunk_transcripts = {}
        # prefer pipeline chunks for full transcript text
        if chunks:
            for ch in chunks:
                cid = ch.get('chunk_id') or ch.get('id') or ch.get('index')
                if cid is None:
                    continue
                txt = (ch.get('text') or '').strip()
                if txt:
                    chunk_transcripts.setdefault(cid, [])
                    if txt not in chunk_transcripts[cid]:
                        chunk_transcripts[cid].append(txt)
        # fallback to highlights
        if not chunk_transcripts and highlights:
            for h in highlights:
                cid = h.get('chunk_id') or h.get('chunk')
                if cid is None:
                    continue
                chunk_transcripts.setdefault(cid, [])
                txt = (h.get('text') or '').strip()
                if txt and txt not in chunk_transcripts[cid]:
                    chunk_transcripts[cid].append(txt)

        # prepare maps for first_cycle, second_cycle, narratives, and memos
        chunk_ids = set()
        fc_map = {}
        if first_cycle:
            for e in first_cycle:
                cid = e.get('chunk_id') or e.get('chunk')
                if cid is not None:
                    chunk_ids.add(cid)
                    fc_map.setdefault(cid, []).append(e.get('code_name') or e.get('code') or e.get('label') or '')
        if chunk_transcripts:
            for cid in chunk_transcripts.keys():
                chunk_ids.add(cid)

        # second-cycle mapping
        narrative_map = {}
        if codeweaving_items:
            for it in codeweaving_items:
                codes = it.get('codes') or []
                textn = it.get('narrative') or (it.get('assertion') and (it.get('assertion').get('assertion_text') or it.get('assertion').get('assertion'))) or ''
                for c in codes:
                    narrative_map.setdefault(c, []).append(textn)

        second_map = {}
        if second_cycle:
            for e in second_cycle:
                # robustly extract a narrative from several possible fields
                narr = ''
                narr = narr or e.get('assertion_text') or ''
                if not narr and isinstance(e.get('assertion'), dict):
                    narr = e.get('assertion').get('assertion_text') or e.get('assertion').get('assertion') or ''
                if not narr:
                    narr = e.get('theme') or e.get('text') or ''
                for cid in e.get('chunks', []) or []:
                    if cid is None:
                        continue
                    if narr:
                        second_map.setdefault(cid, []).append(narr)

        memos_map = {}
        # memos can be simple entries (from memo_engine) or rich memos; support both
        if memos:
            for m in memos:
                cid = m.get('chunk_id') or m.get('chunk')
                textm = (m.get('memo_text') or m.get('text') or '').strip()
                if cid is None:
                    # collect global memos under special key
                    memos_map.setdefault('__global__', []).append(textm)
                    continue
                if not textm:
                    continue
                memos_map.setdefault(cid, []).append(textm)
        
        # Title and original text
        h = doc.add_heading('Qualitative Coding Report', level=1)
        h.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        doc.add_heading('Original Corpus', level=2)
        p = doc.add_paragraph()
        p.add_run(text).font.size = Pt(10)

        # If we have chunk ids, build a table similar to the PDF export
        try:
            chunk_ids = set()
            if first_cycle:
                for e in first_cycle:
                    cid = e.get('chunk_id') or e.get('chunk')
                    if cid is not None:
                        chunk_ids.add(cid)
            if chunk_transcripts:
                for cid in chunk_transcripts.keys():
                    chunk_ids.add(cid)
            if second_cycle:
                for th in second_cycle:
                    for cid in th.get('chunks', []) or []:
                        if cid is not None:
                            chunk_ids.add(cid)

            if chunk_ids:
                doc.add_heading('Analysis Results', level=2)
                # create table with header: No, Transcript, First Cycle, Second Cycle, Memo
                table = doc.add_table(rows=1, cols=5)
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'No'
                hdr_cells[1].text = 'Transcript'
                hdr_cells[2].text = 'First Cycle'
                hdr_cells[3].text = 'Second Cycle'
                hdr_cells[4].text = 'Memo'
                for cid in sorted(chunk_ids):
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(cid)
                    transcript = ''
                    if cid in chunk_transcripts:
                        transcript = '\n'.join(chunk_transcripts[cid])
                    # limit transcript length but keep reasonable amount
                    row_cells[1].text = transcript[:10000]
                    codes = fc_map.get(cid, [])
                    row_cells[2].text = ', '.join([c for c in codes if c])
                    # second cycle narration: prefer second_cycle entries (assertion text or theme), fall back to codeweaving narratives
                    second_narrs = []
                    for n in second_map.get(cid, []):
                        if n and n not in second_narrs:
                            second_narrs.append(n)
                    if not second_narrs:
                        # try narratives from codeweaving items based on codes
                        for c in codes:
                            for n in narrative_map.get(c, []):
                                if n and n not in second_narrs:
                                    second_narrs.append(n)
                    row_cells[3].text = ('\n'.join(second_narrs)) if second_narrs else '-'
                    # memos for this chunk
                    memo_text = '-'
                    if cid in memos_map:
                        memo_text = '\n'.join(memos_map.get(cid, []))
                    row_cells[4].text = memo_text

        except Exception:
            pass

        # Memos
        if memos:
            doc.add_heading('Memos / Analytic Notes', level=2)
            for m in memos:
                p = doc.add_paragraph()
                p.add_run(f"{m.get('created','')} {m.get('author','')}: {m.get('text','')}").font.size = Pt(10)

        doc.save(str(path))
        return str(path)

    def export_pdf(self, doc_id, text, highlights=None, memos=None, module_analyses=None, first_cycle=None, second_cycle=None, codeweaving_items=None, chunks=None, uploaded_filename=None, uploaded_at=None, out_name=None):
        # determine output filename: prefer explicit out_name, then uploaded_filename, then doc_id
        if out_name:
            out_name_final = out_name
        else:
            ext = 'pdf'
            if uploaded_filename:
                name_base = Path(uploaded_filename).stem
            elif doc_id and doc_id != 'default':
                name_base = str(doc_id)
            else:
                name_base = 'export'
            out_name_final = f"Coded - {name_base}.{ext}"
        path = self.storage_dir / out_name_final
        doc = SimpleDocTemplate(str(path), pagesize=letter, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=72)
        styles = getSampleStyleSheet()
        # paragraph styles for wrapping table cells
        wrap_style = ParagraphStyle('wrap', parent=styles['Normal'], fontSize=9, leading=11, wordWrap='CJK')
        header_style = ParagraphStyle('hdr', parent=styles['Normal'], fontSize=10, leading=12)
        flow = []

        # Header metadata
        exported_dt = datetime.datetime.utcnow()
        def fmt_dt(val):
            if not val:
                return 'N/A'
            try:
                if isinstance(val, (int, float)):
                    dt = datetime.datetime.fromtimestamp(val)
                else:
                    # try ISO parse
                    dt = datetime.datetime.fromisoformat(val)
                # include time
                return dt.strftime('%d %B %Y %H:%M:%S')
            except Exception:
                try:
                    # fallback: attempt parse common formats
                    dt = datetime.datetime.strptime(val, '%Y-%m-%d %H:%M:%S')
                    return dt.strftime('%d %B %Y %H:%M:%S')
                except Exception:
                    return str(val)

        exported_label = exported_dt.strftime('%d %B %Y %H:%M:%S')
        file_label = uploaded_filename or 'N/A'
        try:
            if file_label and isinstance(file_label, str):
                file_label = Path(file_label).name
        except Exception:
            pass
        uploaded_label = fmt_dt(uploaded_at)
        generated_by = 'Qualitative Coding Assistant v1.7 Build 002'

        # add report title and metadata block at top
        title = Paragraph('QUALITATIVE ANALYSIS REPORT', styles['Title'])
        flow.append(title)
        flow.append(Spacer(1, 12))
        # metadata lines (left-aligned, normal weight)
        meta_style = ParagraphStyle('meta', parent=styles['Normal'], fontSize=10, leading=12)
        flow.append(Paragraph(f'File: {html.escape(file_label)}', meta_style))
        flow.append(Paragraph(f'Uploaded: {html.escape(uploaded_label)}', meta_style))
        flow.append(Paragraph(f'Exported: {html.escape(exported_label)}', meta_style))
        flow.append(Paragraph(f'Generated by: {html.escape(generated_by)}', meta_style))
        flow.append(Spacer(1, 8))
        # horizontal divider
        flow.append(HRFlowable(width='100%', thickness=1, lineCap='round', color=colors.grey))
        flow.append(Spacer(1, 12))

        # Build chunk-based table only. Collect chunk transcripts preferring pipeline `chunks`, fallback to highlights.
        chunk_transcripts = {}
        # prefer chunks produced by pipeline for full transcripts
        if chunks:
            for ch in chunks:
                cid = ch.get('chunk_id') or ch.get('id') or ch.get('index')
                if cid is None:
                    continue
                txt = (ch.get('text') or '').strip()
                if txt:
                    chunk_transcripts.setdefault(cid, [])
                    if txt not in chunk_transcripts[cid]:
                        chunk_transcripts[cid].append(txt)
        # otherwise, fall back to highlight texts
        if not chunk_transcripts and highlights:
            for h in highlights:
                cid = h.get('chunk_id') or h.get('chunk')
                if cid is None:
                    continue
                chunk_transcripts.setdefault(cid, [])
                txt = (h.get('text') or '').strip()
                if txt and txt not in chunk_transcripts[cid]:
                    chunk_transcripts[cid].append(txt)

        # Build the requested chunk table if we have first_cycle or chunk_transcripts
        try:
            # Collect chunk ids from first_cycle, chunk_transcripts, and second_cycle
            chunk_ids = set()
            fc_map = {}
            if first_cycle:
                for e in first_cycle:
                    cid = e.get('chunk_id') or e.get('chunk')
                    if cid is not None:
                        chunk_ids.add(cid)
                        fc_map.setdefault(cid, []).append(e.get('code_name') or e.get('code') or e.get('label') or '')
            if chunk_transcripts:
                for cid, texts in chunk_transcripts.items(): 
                    chunk_ids.add(cid)
            # also include chunk ids referenced in second_cycle entries
            if second_cycle:
                for th in second_cycle:
                    for cid in th.get('chunks', []) or []:
                        if cid is not None:
                            chunk_ids.add(cid)
            # prepare codeweaving map by codes -> narratives
            narrative_map = {}
            if codeweaving_items:
                for it in codeweaving_items:
                    codes = it.get('codes') or []
                    textn = it.get('narrative') or (it.get('assertion') and (it.get('assertion').get('assertion_text') or it.get('assertion').get('assertion'))) or ''
                    for c in codes:
                        narrative_map.setdefault(c, []).append(textn)

            # prepare second-cycle map: chunk_id -> list of narrations/themes
            second_map = {}
            if second_cycle:
                for e in second_cycle:
                    # robustly extract a narrative from several possible fields
                    narr = ''
                    narr = narr or e.get('assertion_text') or ''
                    if not narr and isinstance(e.get('assertion'), dict):
                        narr = e.get('assertion').get('assertion_text') or e.get('assertion').get('assertion') or ''
                    if not narr:
                        narr = e.get('theme') or e.get('text') or ''
                    for cid in e.get('chunks', []) or []:
                        if cid is None:
                            continue
                        if narr:
                            second_map.setdefault(cid, []).append(narr)

            # prepare memo map: chunk_id -> joined memos
            memos_map = {}
            if memos:
                for m in memos:
                    cid = m.get('chunk_id') or m.get('chunk')
                    txt = (m.get('memo_text') or m.get('text') or '').strip()
                    if cid is None:
                        continue
                    if not txt:
                        continue
                    memos_map.setdefault(cid, []).append(txt)

            if chunk_ids:
                flow.append(Paragraph('Analysis Results', styles['Heading2']))
                # table header (use Paragraphs so ReportLab wraps and measures correctly)
                data = [[Paragraph('<b>No</b>', header_style), Paragraph('<b>Transcript</b>', header_style), Paragraph('<b>First Cycle</b>', header_style), Paragraph('<b>Second Cycle</b>', header_style), Paragraph('<b>Memo</b>', header_style)]]
                for cid in sorted(chunk_ids):
                    transcript = ''
                    if cid in chunk_transcripts:
                        transcript = '\n'.join(chunk_transcripts[cid])
                    # first-cycle codes for this chunk
                    codes = fc_map.get(cid, [])
                    codes_text = ', '.join([c for c in codes if c])
                    # second-cycle: prefer second_cycle entries (assertion text or theme), fall back to codeweaving narratives
                    second_narrs = []
                    for n in second_map.get(cid, []):
                        if n and n not in second_narrs:
                            second_narrs.append(n)
                    if not second_narrs:
                        for c in codes:
                            for n in narrative_map.get(c, []):
                                if n and n not in second_narrs:
                                    second_narrs.append(n)
                    second_text = '\n'.join(second_narrs) if second_narrs else '-'
                    # memo for this chunk (if any)
                    memo_text = '-'
                    if cid in memos_map:
                        memo_text = '\n'.join(memos_map.get(cid, []))
                    data.append([
                        Paragraph(str(cid), wrap_style),
                        Paragraph((transcript[:1000] or '-').replace('\n', '<br/>'), wrap_style),
                        Paragraph(codes_text or '-', wrap_style),
                        Paragraph(second_text.replace('\n', '<br/>'), wrap_style),
                        Paragraph(memo_text.replace('\n', '<br/>') if isinstance(memo_text, str) else '-', wrap_style)
                    ])
                # layout table — explicit widths that fit printable area (6.5in)
                # columns: No, Transcript, First Cycle, Second Cycle, Memo
                col_widths = [0.6*inch, 2.8*inch, 1.0*inch, 1.1*inch, 1.0*inch]
                t = Table(data, colWidths=col_widths)
                t.setStyle(TableStyle([('GRID',(0,0),(-1,-1),0.5,colors.grey),('BACKGROUND',(0,0),(-1,0),colors.lightgrey),('VALIGN',(0,0),(-1,-1),'TOP')] ))
                flow.append(t)
                flow.append(Spacer(1,12))
        except Exception:
            # swallow export table errors and continue
            pass

        # Include any global memos (not linked to specific chunks)
        try:
            if memos_map and isinstance(memos_map.get('__global__'), list) and len(memos_map.get('__global__')):
                flow.append(Paragraph('Memos / Analytic Notes', styles['Heading2']))
                for gm in memos_map.get('__global__'):
                    flow.append(Paragraph(html.escape(gm), wrap_style))
                flow.append(Spacer(1,12))
        except Exception:
            pass

        # do not include coding summary or module analyses in PDF — only the chunk table

        doc.build(flow)
        return str(path)
